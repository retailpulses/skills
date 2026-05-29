import os
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def create_payslip_docx(payload, month_label, output_path):
    doc = Document()
    
    # Set Japanese font (MS Mincho or similar)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'MS Mincho'
    font.size = Pt(10)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), 'MS Mincho')

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"給与明細書 ({month_label})")
    run.bold = True
    run.size = Pt(16)

    # Employee Name and Date
    header = doc.add_table(rows=1, cols=2)
    header.width = Inches(6)
    cells = header.rows[0].cells
    cells[0].text = f"氏名: {payload['従業員名']} 様"
    cells[1].text = f"支給日: {payload['給与月']}"
    cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

    doc.add_paragraph()

    # Earnings Table
    doc.add_heading('支給項目 (Earnings)', level=2)
    earnings_table = doc.add_table(rows=2, cols=3)
    earnings_table.style = 'Table Grid'
    hdr_cells = earnings_table.rows[0].cells
    hdr_cells[0].text = '基本給'
    hdr_cells[1].text = 'その他手当'
    hdr_cells[2].text = '総支給額'
    
    row_cells = earnings_table.rows[1].cells
    row_cells[0].text = f"¥{int(payload['基本給']):,}"
    row_cells[1].text = f"¥{int(payload['その他手当']):,}"
    gross = int(payload['基本給']) + int(payload['その他手当'])
    row_cells[2].text = f"¥{gross:,}"

    doc.add_paragraph()

    # Deductions Table
    doc.add_heading('控除項目 (Deductions)', level=2)
    deductions_table = doc.add_table(rows=2, cols=6)
    deductions_table.style = 'Table Grid'
    hdr_cells = deductions_table.rows[0].cells
    hdr_cells[0].text = '健康保険'
    hdr_cells[1].text = '厚生年金'
    hdr_cells[2].text = '雇用保険'
    hdr_cells[3].text = '所得税'
    hdr_cells[4].text = '住民税'
    hdr_cells[5].text = '控除計'
    
    row_cells = deductions_table.rows[1].cells
    row_cells[0].text = f"¥{int(payload['健康保険料']):,}"
    row_cells[1].text = f"¥{int(payload['厚生年金保険料']):,}"
    row_cells[2].text = f"¥{int(payload['雇用保険料']):,}"
    row_cells[3].text = f"¥{int(payload['所得税']):,}"
    row_cells[4].text = f"¥{int(payload['住民税']):,}"
    total_deductions = int(payload['健康保険料']) + int(payload['厚生年金保険料']) + int(payload['雇用保険料']) + int(payload['所得税']) + int(payload['住民税'])
    row_cells[5].text = f"¥{total_deductions:,}"

    doc.add_paragraph()

    # Payout Table
    doc.add_heading('振込額 (Payout)', level=2)
    payout_table = doc.add_table(rows=2, cols=3)
    payout_table.style = 'Table Grid'
    hdr_cells = payout_table.rows[0].cells
    hdr_cells[0].text = '差引支給額 (Net)'
    hdr_cells[1].text = '振込額 (Actual)'
    hdr_cells[2].text = '振込方法'
    
    row_cells = payout_table.rows[1].cells
    net_pay = gross - total_deductions
    row_cells[0].text = f"¥{net_pay:,}"
    row_cells[1].text = f"¥{int(payload['実支給額']):,}"
    row_cells[2].text = payload.get('振込方法', '銀行振り込み')

    doc.add_paragraph()

    # Notes
    doc.add_heading('備考 (Remarks)', level=2)
    doc.add_paragraph(payload.get('Notes', ''))

    doc.save(output_path)
    return output_path

if __name__ == "__main__":
    # Test
    sample_payload = {
        "従業員名": "楊　永亮",
        "給与月": "2026-05-31",
        "基本給": 400000,
        "その他手当": 0,
        "健康保険料": 20664,
        "厚生年金保険料": 37515,
        "雇用保険料": 0,
        "所得税": 10750,
        "住民税": 0,
        "実支給額": 331000,
        "振込方法": "銀行振り込み",
        "Notes": "差引支給額と実支給額（固定振込額）との差額については、12月の年末調整時または退職時に一括して精算するものとします。"
    }
    create_payslip_docx(sample_payload, "2026年5月", "test_payslip.docx")
    print("Test docx created.")
